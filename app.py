# =============================================================================
# COLEGIO PUKARAY — Sistema de Gestión de Entrevistas de Convivencia Escolar
# Versión: 2.0.0 | Python 3.11+ | Streamlit
# =============================================================================
# Autor: Arquitecto de Software Institucional
# Descripción: Plataforma interna para gestión, registro, exportación y
#              revisión de entrevistas de convivencia escolar.
# =============================================================================

import streamlit as st
import pandas as pd
import sqlite3
import hashlib
import shutil
import os
import io
import base64
from pathlib import Path
from datetime import datetime, date, time
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURACIÓN GLOBAL DE PÁGINA
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Pukaray | Convivencia Escolar",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────────────────────────────────────
# CSS INSTITUCIONAL — Paleta: Verde oscuro, Verde medio, Burdeo, Rojo, Crema
# ─────────────────────────────────────────────────────────────────────────────

INSTITUTIONAL_CSS = """
<style>
    /* ── Variables de color institucional ── */
    :root {
        --verde-oscuro:  #1B4332;
        --verde-medio:   #2D6A4F;
        --verde-claro:   #40916C;
        --burdeo:        #6B2737;
        --rojo:          #C0392B;
        --crema:         #FAF3E0;
        --crema-oscuro:  #EDE0C4;
        --texto-oscuro:  #1A1A1A;
        --gris-claro:    #F0F0F0;
        --blanco:        #FFFFFF;
    }

    /* ── Fondo general ── */
    .main, .block-container {
        background-color: var(--crema) !important;
        color: var(--texto-oscuro) !important;
        font-family: 'Georgia', 'Times New Roman', serif;
    }

    /* ── Sidebar ── */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, var(--verde-oscuro) 0%, var(--verde-medio) 100%) !important;
        border-right: 3px solid var(--burdeo);
    }
    [data-testid="stSidebar"] * {
        color: var(--crema) !important;
        font-family: 'Georgia', serif !important;
    }
    [data-testid="stSidebar"] .stRadio label {
        font-size: 0.95rem !important;
        padding: 4px 0 !important;
    }
    [data-testid="stSidebar"] hr {
        border-color: var(--crema-oscuro) !important;
        opacity: 0.4;
    }

    /* ── Encabezado principal ── */
    .header-institucional {
        background: var(--verde-oscuro);
        color: var(--crema) !important;
        padding: 1.2rem 2rem;
        border-radius: 8px;
        border-left: 6px solid var(--burdeo);
        margin-bottom: 1.5rem;
        display: flex;
        align-items: center;
        gap: 1rem;
    }
    .header-institucional h1 {
        color: var(--crema) !important;
        font-size: 1.6rem;
        margin: 0;
        font-weight: 700;
        letter-spacing: 0.02em;
    }
    .header-institucional p {
        color: var(--crema-oscuro) !important;
        margin: 0;
        font-size: 0.85rem;
    }

    /* ── Tarjetas de sección ── */
    .seccion-card {
        background: var(--blanco);
        border: 1px solid var(--crema-oscuro);
        border-top: 4px solid var(--verde-medio);
        border-radius: 6px;
        padding: 1.2rem 1.5rem;
        margin-bottom: 1.2rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
    }
    .seccion-card h3 {
        color: var(--verde-oscuro) !important;
        font-size: 1.05rem;
        margin-bottom: 0.8rem;
        border-bottom: 1px solid var(--crema-oscuro);
        padding-bottom: 0.4rem;
    }

    /* ── Alerta de respaldo ── */
    .alerta-respaldo {
        background: #FFF3CD;
        border: 1px solid #856404;
        border-left: 5px solid #856404;
        border-radius: 4px;
        padding: 0.7rem 1rem;
        margin: 0.5rem 0;
        font-size: 0.88rem;
        color: #533F03;
    }

    /* ── Badge de estado ── */
    .badge-abierto     { background:#2D6A4F; color:#fff; padding:3px 10px; border-radius:12px; font-size:0.8rem; font-weight:600; }
    .badge-proceso     { background:#F4A261; color:#1A1A1A; padding:3px 10px; border-radius:12px; font-size:0.8rem; font-weight:600; }
    .badge-seguimiento { background:#2196F3; color:#fff; padding:3px 10px; border-radius:12px; font-size:0.8rem; font-weight:600; }
    .badge-cerrado     { background:#6B2737; color:#fff; padding:3px 10px; border-radius:12px; font-size:0.8rem; font-weight:600; }
    .badge-noaplica    { background:#888; color:#fff; padding:3px 10px; border-radius:12px; font-size:0.8rem; font-weight:600; }

    /* ── Botones personalizados ── */
    .stButton > button {
        background: var(--verde-medio) !important;
        color: var(--crema) !important;
        border: none !important;
        border-radius: 5px !important;
        font-weight: 600 !important;
        font-family: 'Georgia', serif !important;
        transition: background 0.2s;
    }
    .stButton > button:hover {
        background: var(--verde-oscuro) !important;
        color: var(--crema) !important;
    }
    .btn-danger > button {
        background: var(--rojo) !important;
    }
    .btn-burdeo > button {
        background: var(--burdeo) !important;
    }

    /* ── Inputs y select ── */
    .stTextInput input, .stSelectbox select, .stTextArea textarea,
    .stDateInput input, .stTimeInput input, .stNumberInput input {
        border: 1px solid var(--crema-oscuro) !important;
        border-radius: 4px !important;
        background: var(--blanco) !important;
        font-family: 'Georgia', serif !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--verde-medio) !important;
        box-shadow: 0 0 0 2px rgba(45,106,79,0.2) !important;
    }

    /* ── Tablas ── */
    .dataframe thead tr th {
        background: var(--verde-oscuro) !important;
        color: var(--crema) !important;
    }

    /* ── Login card ── */
    .login-card {
        max-width: 420px;
        margin: 4rem auto;
        background: var(--blanco);
        border-radius: 10px;
        padding: 2.5rem;
        box-shadow: 0 8px 32px rgba(27,67,50,0.15);
        border-top: 6px solid var(--burdeo);
    }
    .login-card h2 {
        color: var(--verde-oscuro) !important;
        text-align: center;
        margin-bottom: 0.3rem;
    }
    .login-card p {
        text-align: center;
        color: #666;
        font-size: 0.88rem;
        margin-bottom: 1.5rem;
    }

    /* ── Timeline CC ── */
    .timeline-item {
        display: flex;
        gap: 12px;
        margin-bottom: 10px;
        align-items: flex-start;
        padding: 8px 10px;
        background: var(--blanco);
        border-radius: 5px;
        border-left: 3px solid var(--verde-medio);
        font-size: 0.9rem;
    }
    .timeline-fecha {
        color: var(--verde-oscuro);
        font-weight: 600;
        min-width: 90px;
        font-size: 0.82rem;
    }

    /* ── Scrollbar ── */
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-track { background: var(--crema-oscuro); }
    ::-webkit-scrollbar-thumb { background: var(--verde-medio); border-radius: 3px; }

    /* ── Ocultar menú y footer de Streamlit ── */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
</style>
"""
st.markdown(INSTITUTIONAL_CSS, unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# RUTAS Y DIRECTORIOS
# ─────────────────────────────────────────────────────────────────────────────

BASE_DIR         = Path(__file__).parent
DATOS_DIR        = BASE_DIR / "datos"
RESPALDOS_DIR    = DATOS_DIR / "respaldos_registro_entrevistas"
PLANTILLAS_DIR   = BASE_DIR / "plantillas"
ASSETS_DIR       = BASE_DIR / "assets"
EXPORTADOS_DIR   = BASE_DIR / "exportados"

EXCEL_REGISTRO   = DATOS_DIR / "registro_entrevistas.xlsx"
EXCEL_BD         = DATOS_DIR / "base_datos_pukaray.xlsx"
DB_PATH          = DATOS_DIR / "usuarios_pukaray.db"

LOGO_PATH        = ASSETS_DIR / "logo_pukaray.png"

# Crear directorios si no existen
for d in [DATOS_DIR, RESPALDOS_DIR, PLANTILLAS_DIR, ASSETS_DIR, EXPORTADOS_DIR]:
    d.mkdir(parents=True, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# COLUMNAS DEL EXCEL DE REGISTRO
# ─────────────────────────────────────────────────────────────────────────────

COLUMNAS_REGISTRO = [
    "id_registro", "tipo_ficha", "fecha", "hora", "depto_cita",
    "curso", "estudiante", "run_estudiante", "participantes",
    "motivo", "tipo_falta_rice", "protocolos", "cc",
    "folio", "n_entrevista", "estado_caso", "condicion_caso",
    "antecedentes", "acuerdos", "compromisos", "observaciones",
    "resumen_libro_clases",
    "check_notificacion", "check_firma_apoderado", "check_firma_estudiante",
    "check_firma_docente", "check_derivacion", "check_seguimiento",
    "check_cierre_cc", "check_copia_apoderado",
    "usuario_registro", "timestamp_registro"
]


# ─────────────────────────────────────────────────────────────────────────────
# UTILIDADES: HASH SHA256
# ─────────────────────────────────────────────────────────────────────────────

def hash_sha256(texto: str) -> str:
    """
    Genera el hash SHA256 de una cadena de texto.
    Utilizado para almacenar contraseñas de forma segura en la base de datos SQLite.
    NUNCA se almacena la contraseña en texto plano.
    """
    return hashlib.sha256(texto.encode("utf-8")).digest().hex()


# ─────────────────────────────────────────────────────────────────────────────
# BASE DE DATOS SQLITE — USUARIOS
# ─────────────────────────────────────────────────────────────────────────────

def inicializar_db():
    """
    Crea la tabla de usuarios en SQLite si no existe.
    Inserta usuarios por defecto: admin / usuario.
    Las contraseñas se almacenan como hashes SHA256.
    """
    conn = sqlite3.connect(str(DB_PATH))
    cur  = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS usuarios (
            id        INTEGER PRIMARY KEY AUTOINCREMENT,
            username  TEXT    UNIQUE NOT NULL,
            password  TEXT    NOT NULL,
            rol       TEXT    NOT NULL DEFAULT 'usuario',
            nombre    TEXT    NOT NULL DEFAULT ''
        )
    """)
    # Usuarios por defecto (solo si no existen)
    defaults = [
        ("admin",    hash_sha256("Admin2024!"),   "administrador", "Administrador Pukaray"),
        ("orientador", hash_sha256("Pukaray2024!"), "usuario",      "Orientador(a) Escolar"),
    ]
    for u in defaults:
        cur.execute(
            "INSERT OR IGNORE INTO usuarios (username, password, rol, nombre) VALUES (?,?,?,?)",
            u
        )
    conn.commit()
    conn.close()


def autenticar_usuario(username: str, password: str):
    """
    Valida credenciales contra la base de datos SQLite.
    Retorna dict con datos del usuario o None si falla.
    """
    try:
        conn = sqlite3.connect(str(DB_PATH))
        cur  = conn.cursor()
        cur.execute(
            "SELECT id, username, rol, nombre FROM usuarios WHERE username=? AND password=?",
            (username.strip(), hash_sha256(password))
        )
        row = cur.fetchone()
        conn.close()
        if row:
            return {"id": row[0], "username": row[1], "rol": row[2], "nombre": row[3]}
        return None
    except Exception as e:
        st.error(f"Error de autenticación: {e}")
        return None


def obtener_todos_los_usuarios():
    """Retorna lista de todos los usuarios (sin contraseña)."""
    try:
        conn = sqlite3.connect(str(DB_PATH))
        df   = pd.read_sql_query(
            "SELECT id, username, rol, nombre FROM usuarios", conn
        )
        conn.close()
        return df
    except Exception:
        return pd.DataFrame(columns=["id", "username", "rol", "nombre"])


def crear_usuario(username, password, rol, nombre):
    """Inserta un nuevo usuario en la base de datos."""
    try:
        conn = sqlite3.connect(str(DB_PATH))
        cur  = conn.cursor()
        cur.execute(
            "INSERT INTO usuarios (username, password, rol, nombre) VALUES (?,?,?,?)",
            (username.strip(), hash_sha256(password), rol, nombre)
        )
        conn.commit()
        conn.close()
        return True, "Usuario creado correctamente."
    except sqlite3.IntegrityError:
        return False, "El nombre de usuario ya existe."
    except Exception as e:
        return False, str(e)


def eliminar_usuario(user_id):
    """Elimina un usuario por ID."""
    try:
        conn = sqlite3.connect(str(DB_PATH))
        cur  = conn.cursor()
        cur.execute("DELETE FROM usuarios WHERE id=?", (user_id,))
        conn.commit()
        conn.close()
        return True
    except Exception:
        return False


# ─────────────────────────────────────────────────────────────────────────────
# RUTINA DE RESPALDO AUTOMÁTICO — REGLA INQUEBRANTABLE
# ─────────────────────────────────────────────────────────────────────────────

def crear_respaldo_excel() -> str:
    """
    Crea una copia de seguridad del archivo registro_entrevistas.xlsx
    ANTES de cualquier operación de escritura, modificación o borrado.
    
    Nomenclatura: registro_entrevistas_antes_de_guardar_AAAAMMDD_HHMMSS.xlsx
    
    Retorna la ruta del archivo de respaldo generado.
    NUNCA debe ser llamado sin verificar que el archivo fuente existe.
    """
    if not EXCEL_REGISTRO.exists():
        # No hay nada que respaldar aún; es la primera escritura
        return ""
    
    timestamp    = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_bkp   = f"registro_entrevistas_antes_de_guardar_{timestamp}.xlsx"
    ruta_bkp     = RESPALDOS_DIR / nombre_bkp
    
    try:
        shutil.copy2(str(EXCEL_REGISTRO), str(ruta_bkp))
        return str(ruta_bkp)
    except Exception as e:
        st.warning(f"⚠️ No se pudo crear respaldo: {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# LECTURA Y ESCRITURA DEL EXCEL DE REGISTRO
# ─────────────────────────────────────────────────────────────────────────────

def cargar_registro() -> pd.DataFrame:
    """
    Lee el archivo registro_entrevistas.xlsx y retorna un DataFrame.
    Si el archivo no existe, retorna un DataFrame vacío con las columnas correctas.
    Captura PermissionError si el archivo está abierto por otro proceso.
    """
    if not EXCEL_REGISTRO.exists():
        return pd.DataFrame(columns=COLUMNAS_REGISTRO)
    try:
        df = pd.read_excel(str(EXCEL_REGISTRO), dtype=str)
        # Garantizar que todas las columnas esperadas existan
        for col in COLUMNAS_REGISTRO:
            if col not in df.columns:
                df[col] = ""
        return df[COLUMNAS_REGISTRO]
    except PermissionError:
        st.error(
            "❌ El archivo `registro_entrevistas.xlsx` está abierto en otro programa. "
            "Ciérrelo e intente nuevamente."
        )
        return pd.DataFrame(columns=COLUMNAS_REGISTRO)
    except Exception as e:
        st.error(f"Error al cargar el registro: {e}")
        return pd.DataFrame(columns=COLUMNAS_REGISTRO)


def guardar_registro(df: pd.DataFrame) -> bool:
    """
    Guarda el DataFrame en registro_entrevistas.xlsx.
    
    Flujo obligatorio:
      1. Crear respaldo automático del estado anterior.
      2. Intentar escritura.
      3. Capturar PermissionError y otros errores.
    
    Retorna True si la escritura fue exitosa.
    """
    # PASO 1: RESPALDO AUTOMÁTICO (regla inquebrantable)
    ruta_bkp = crear_respaldo_excel()
    if ruta_bkp:
        st.markdown(
            f'<div class="alerta-respaldo">🔒 Respaldo creado automáticamente: '
            f'<code>{Path(ruta_bkp).name}</code></div>',
            unsafe_allow_html=True
        )
    
    # PASO 2: ESCRITURA
    try:
        df.to_excel(str(EXCEL_REGISTRO), index=False, engine="openpyxl")
        return True
    except PermissionError:
        st.error(
            "❌ No se pudo guardar: `registro_entrevistas.xlsx` está abierto "
            "por otro programa. Ciérrelo e intente nuevamente."
        )
        return False
    except Exception as e:
        st.error(f"❌ Error al guardar el registro: {e}")
        return False


def generar_id_registro(df: pd.DataFrame) -> str:
    """Genera un ID único correlativo para el nuevo registro."""
    if df.empty or "id_registro" not in df.columns:
        return "ENT-001"
    ids = df["id_registro"].dropna().tolist()
    nums = []
    for i in ids:
        try:
            nums.append(int(str(i).replace("ENT-", "")))
        except ValueError:
            pass
    siguiente = max(nums) + 1 if nums else 1
    return f"ENT-{siguiente:03d}"


# ─────────────────────────────────────────────────────────────────────────────
# GENERADOR DE CÓDIGO CC
# ─────────────────────────────────────────────────────────────────────────────

def generar_cc_nuevo(df: pd.DataFrame, curso: str, anio: int = None) -> str:
    """
    Genera un código de caso (CC) único con formato: CC-CURSO-AAAA-NNN
    Ejemplo: CC-7A-2024-001
    """
    if anio is None:
        anio = datetime.now().year
    curso_clean = curso.replace(" ", "").upper()[:4]
    prefijo     = f"CC-{curso_clean}-{anio}-"
    
    if df.empty or "cc" not in df.columns:
        return f"{prefijo}001"
    
    ccs_del_curso = [c for c in df["cc"].dropna().tolist() if str(c).startswith(prefijo)]
    nums = []
    for c in ccs_del_curso:
        try:
            nums.append(int(str(c).replace(prefijo, "")))
        except ValueError:
            pass
    siguiente = max(nums) + 1 if nums else 1
    return f"{prefijo}{siguiente:03d}"


def obtener_ccs_existentes(df: pd.DataFrame) -> list:
    """Retorna lista de CCs únicos existentes en el registro."""
    if df.empty or "cc" not in df.columns:
        return []
    return sorted(df["cc"].dropna().unique().tolist())


# ─────────────────────────────────────────────────────────────────────────────
# GENERADOR DE DOCUMENTOS WORD (ACTAS)
# ─────────────────────────────────────────────────────────────────────────────

def _reemplazar_marcadores_doc(doc: DocxDocument, datos: dict):
    """
    Recorre todos los párrafos y celdas de tabla del documento Word
    y reemplaza los marcadores {{campo}} con los valores del dict datos.
    Preserva el formato (bold, size, color) del texto original.
    """
    def reemplazar_en_parrafo(p):
        texto_original = p.text
        texto_nuevo    = texto_original
        for clave, valor in datos.items():
            marcador    = "{{" + clave + "}}"
            texto_nuevo = texto_nuevo.replace(marcador, str(valor) if valor else "")
        
        if texto_nuevo != texto_original:
            # Preservar formato del primer run
            fmt = {}
            if p.runs:
                r0       = p.runs[0]
                fmt_run  = r0.font
                fmt      = {
                    "bold":      fmt_run.bold,
                    "italic":    fmt_run.italic,
                    "size":      fmt_run.size,
                    "color":     fmt_run.color.rgb if fmt_run.color and fmt_run.color.type else None,
                    "name":      fmt_run.name,
                }
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = texto_nuevo
                if fmt.get("bold")   is not None: p.runs[0].font.bold   = fmt["bold"]
                if fmt.get("italic") is not None: p.runs[0].font.italic = fmt["italic"]
                if fmt.get("size"):               p.runs[0].font.size   = fmt["size"]
                if fmt.get("name"):               p.runs[0].font.name   = fmt["name"]
                if fmt.get("color"):
                    try:
                        p.runs[0].font.color.rgb = fmt["color"]
                    except Exception:
                        pass
            else:
                p.add_run(texto_nuevo)

    # Párrafos del cuerpo
    for p in doc.paragraphs:
        reemplazar_en_parrafo(p)
    # Celdas de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    reemplazar_en_parrafo(p)


def _crear_doc_desde_cero(datos: dict, tipo: str) -> DocxDocument:
    """
    Genera un documento Word institucional desde cero cuando no existe
    la plantilla en /plantillas/. Usa python-docx con formato completo.
    """
    doc = DocxDocument()
    
    # ── Estilos base ──
    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)
    
    # ── Encabezado del documento ──
    enc = doc.add_heading("", level=0)
    enc.clear()
    run_enc = enc.add_run(f"COLEGIO PUKARAY")
    run_enc.font.bold      = True
    run_enc.font.size      = Pt(16)
    run_enc.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)  # Verde oscuro
    enc.alignment          = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_enc = doc.add_paragraph(
        "Departamento de Convivencia Escolar — Entrevistas de Convivencia"
    )
    sub_enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_enc.runs[0].font.size      = Pt(11)
    sub_enc.runs[0].font.color.rgb = RGBColor(0x6B, 0x27, 0x37)  # Burdeo
    
    doc.add_paragraph("")
    
    # ── Título del acta ──
    titulo = "ACTA DE ENTREVISTA DE CONVIVENCIA ESCOLAR"
    if tipo == "apoderado":
        titulo += " — APODERADO/FAMILIA"
    else:
        titulo += " — ESTUDIANTE"
    
    t = doc.add_heading(titulo, level=1)
    t.runs[0].font.color.rgb = RGBColor(0x1B, 0x43, 0x32)
    t.alignment              = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("")
    
    # ── Tabla de datos generales ──
    campos_generales = [
        ("Código de Caso (CC)", datos.get("cc", "")),
        ("Folio",               datos.get("folio", "")),
        ("N° Entrevista",       datos.get("n_entrevista", "")),
        ("Fecha",               datos.get("fecha", "")),
        ("Hora",                datos.get("hora", "")),
        ("Departamento que cita", datos.get("depto_cita", "")),
        ("Curso",               datos.get("curso", "")),
        ("Estudiante",          datos.get("estudiante", "")),
        ("RUN Estudiante",      datos.get("run_estudiante", "")),
        ("Participantes",       datos.get("participantes", "")),
        ("Estado del Caso",     datos.get("estado_caso", "")),
        ("Condición del Caso",  datos.get("condicion_caso", "")),
        ("Tipo Falta RICE",     datos.get("tipo_falta_rice", "")),
        ("Protocolos",          datos.get("protocolos", "")),
        ("Motivo",              datos.get("motivo", "")),
    ]
    
    tabla = doc.add_table(rows=len(campos_generales), cols=2)
    tabla.style = "Table Grid"
    
    for i, (label, valor) in enumerate(campos_generales):
        celda_l = tabla.cell(i, 0)
        celda_r = tabla.cell(i, 1)
        
        p_l = celda_l.paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.font.bold      = True
        r_l.font.size      = Pt(10)
        r_l.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)
        
        p_r = celda_r.paragraphs[0]
        r_r = p_r.add_run(str(valor))
        r_r.font.size = Pt(10)
        
        # Ancho de columnas
        celda_l.width = Inches(2.4)
        celda_r.width = Inches(4.1)
    
    doc.add_paragraph("")
    
    # ── Secciones de texto extendido ──
    secciones_texto = [
        ("ANTECEDENTES DEL CASO",              datos.get("antecedentes", "")),
        ("ACUERDOS ALCANZADOS",                datos.get("acuerdos", "")),
        ("COMPROMISOS ADQUIRIDOS",             datos.get("compromisos", "")),
        ("OBSERVACIONES",                      datos.get("observaciones", "")),
        ("RESUMEN LIBRO DE CLASES",            datos.get("resumen_libro_clases", "")),
    ]
    
    for titulo_sec, contenido in secciones_texto:
        h = doc.add_heading(titulo_sec, level=2)
        h.runs[0].font.color.rgb = RGBColor(0x6B, 0x27, 0x37)  # Burdeo
        h.runs[0].font.size      = Pt(12)
        cuerpo = doc.add_paragraph(contenido if contenido else "— Sin registro —")
        cuerpo.runs[0].font.size = Pt(10)
        doc.add_paragraph("")
    
    # ── Checklist de cierre ──
    h_check = doc.add_heading("CHECKLIST DE CIERRE", level=2)
    h_check.runs[0].font.color.rgb = RGBColor(0x6B, 0x27, 0x37)
    
    checklist_items = [
        ("Notificación realizada",          datos.get("check_notificacion", "")),
        ("Firma apoderado",                 datos.get("check_firma_apoderado", "")),
        ("Firma estudiante",                datos.get("check_firma_estudiante", "")),
        ("Firma docente/encargado",         datos.get("check_firma_docente", "")),
        ("Derivación realizada",            datos.get("check_derivacion", "")),
        ("Seguimiento programado",          datos.get("check_seguimiento", "")),
        ("Cierre de CC",                    datos.get("check_cierre_cc", "")),
        ("Copia entregada a apoderado",     datos.get("check_copia_apoderado", "")),
    ]
    
    for item, val in checklist_items:
        simbolo = "☑" if str(val).lower() in ("true", "si", "sí", "1", "yes") else "☐"
        p = doc.add_paragraph(f"{simbolo}  {item}")
        p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph("")
    
    # ── Firmas ──
    doc.add_heading("FIRMAS", level=2)
    firmas = doc.add_table(rows=3, cols=3)
    firmas.style = "Table Grid"
    
    encabezados_firma = ["Apoderado / Familiar", "Estudiante", "Encargado Convivencia"]
    for j, ef in enumerate(encabezados_firma):
        c = firmas.cell(0, j)
        r = c.paragraphs[0].add_run(ef)
        r.font.bold = True
        r.font.size = Pt(9)
    
    for j in range(3):
        firmas.cell(1, j).paragraphs[0].add_run("\n\n\n")
        linea = firmas.cell(2, j).paragraphs[0]
        linea.add_run("_______________________").font.size = Pt(9)
    
    doc.add_paragraph("")
    
    # ── Pie de página ──
    pie = doc.add_paragraph(
        f"Documento generado: {datetime.now().strftime('%d/%m/%Y %H:%M')} | "
        f"Sistema Pukaray Convivencia Escolar | Registro: {datos.get('id_registro', '')}"
    )
    pie.runs[0].font.size = Pt(8)
    pie.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def generar_word(datos: dict, tipo: str = "apoderado") -> bytes:
    """
    Genera el documento Word del acta de entrevista.
    
    Flujo:
    1. Intentar cargar la plantilla correspondiente desde /plantillas/.
    2. Si la plantilla existe, reemplazar marcadores {{campo}}.
    3. Si no existe, generar el documento desde cero con python-docx.
    4. Retornar los bytes del archivo para descarga en Streamlit.
    
    tipo: "apoderado" | "estudiante"
    """
    nombre_plantilla = (
        "Ficha_Entrevista_APODERADO.docx"
        if tipo == "apoderado"
        else "Ficha_Entrevista_ESTUDIANTE.docx"
    )
    ruta_plantilla = PLANTILLAS_DIR / nombre_plantilla
    
    try:
        if ruta_plantilla.exists():
            doc = DocxDocument(str(ruta_plantilla))
            _reemplazar_marcadores_doc(doc, datos)
        else:
            doc = _crear_doc_desde_cero(datos, tipo)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    except Exception as e:
        st.error(f"Error generando documento Word: {e}")
        return b""


# ─────────────────────────────────────────────────────────────────────────────
# INICIALIZACIÓN DE SESSION_STATE
# ─────────────────────────────────────────────────────────────────────────────

def init_session_state():
    """
    Inicializa todas las variables de session_state al arrancar la app.
    Esto garantiza que no haya mutación directa de widgets ya instanciados.
    Se sigue el patrón: definir variables auxiliares en session_state y
    usar callbacks (on_click, on_change) para modificarlas.
    """
    defaults = {
        # Autenticación
        "autenticado":         False,
        "usuario":             None,
        # Navegación
        "pagina_actual":       "login",
        # Estado formularios (se limpian explícitamente con callbacks)
        "form_guardado":       False,
        "form_tipo":           "apoderado",
        "form_cc_modo":        "nuevo",
        "form_cc_seleccionado": "",
        # Precarga de datos de CC existente en formulario
        # Cuando el usuario elige un CC existente, se carga el último registro
        # de ese CC para precompletar los campos del formulario.
        "form_preload":        None,   # dict con datos precargados o None
        "form_cc_prev":        "",     # CC seleccionado previamente (para detectar cambio)
        # Datos en caché
        "df_registro":         None,
        "df_cargado_en":       None,
        # Vista de revisión
        "revision_idx":        None,
        # Borrado seguro
        "borrado_confirmado":  False,
        "borrado_indices":     [],
        # Administración
        "admin_tab":           0,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ─────────────────────────────────────────────────────────────────────────────
# CALLBACKS
# ─────────────────────────────────────────────────────────────────────────────

def cb_logout():
    """Callback para cerrar sesión."""
    st.session_state["autenticado"]   = False
    st.session_state["usuario"]       = None
    st.session_state["pagina_actual"] = "login"


def cb_limpiar_form():
    """Callback para limpiar el formulario de registro."""
    st.session_state["form_guardado"] = False


def cb_cc_modo(modo: str):
    """Callback para cambiar el modo de asignación de CC."""
    st.session_state["form_cc_modo"] = modo


def cb_login(username: str, password: str):
    """Callback para iniciar sesión."""
    usuario = autenticar_usuario(username, password)
    if usuario:
        st.session_state["autenticado"]   = True
        st.session_state["usuario"]       = usuario
        st.session_state["pagina_actual"] = "formulario"
    else:
        st.session_state["_login_error"] = True


def recargar_df():
    """Fuerza la recarga del DataFrame desde disco."""
    st.session_state["df_registro"]   = cargar_registro()
    st.session_state["df_cargado_en"] = datetime.now()


def obtener_df() -> pd.DataFrame:
    """
    Retorna el DataFrame del registro.
    Usa caché en session_state para no releer disco en cada re-render,
    a menos que se fuerce la recarga.
    """
    if st.session_state["df_registro"] is None:
        recargar_df()
    return st.session_state["df_registro"].copy()


# ─────────────────────────────────────────────────────────────────────────────
# COMPONENTES UI REUTILIZABLES
# ─────────────────────────────────────────────────────────────────────────────

def mostrar_header(titulo: str, subtitulo: str = ""):
    """Muestra el encabezado institucional de cada sección."""
    logo_html = ""
    if LOGO_PATH.exists():
        with open(str(LOGO_PATH), "rb") as f:
            logo_b64 = base64.b64encode(f.read()).decode()
        logo_html = (
            f'<img src="data:image/png;base64,{logo_b64}" '
            f'style="height:52px;border-radius:4px;" />'
        )
    st.markdown(
        f"""
        <div class="header-institucional">
            {logo_html}
            <div>
                <h1>🏫 {titulo}</h1>
                <p>Colegio Pukaray — Departamento de Convivencia Escolar{" | " + subtitulo if subtitulo else ""}</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def badge_estado(estado: str) -> str:
    """Retorna HTML de badge según el estado del caso."""
    mapa = {
        "abierto":     ("badge-abierto",     "🟢 Abierto"),
        "en proceso":  ("badge-proceso",     "🟡 En Proceso"),
        "seguimiento": ("badge-seguimiento", "🔵 Seguimiento"),
        "cerrado":     ("badge-cerrado",     "🔴 Cerrado"),
        "no aplica":   ("badge-noaplica",    "⚪ No Aplica"),
    }
    clase, texto = mapa.get(str(estado).lower(), ("badge-noaplica", estado))
    return f'<span class="{clase}">{texto}</span>'


# ─────────────────────────────────────────────────────────────────────────────
# MÓDULO 1: LOGIN INSTITUCIONAL
# ─────────────────────────────────────────────────────────────────────────────

def pagina_login():
    """
    Módulo de autenticación institucional.
    - Validación SHA256 contra SQLite.
    - Campos: usuario, contraseña (type='password').
    - Roles: Administrador y Usuario.
    - NO se almacenan contraseñas en texto plano.
    """
    # Centrar el login card
    col_l, col_c, col_r = st.columns([1, 1.4, 1])
    with col_c:
        st.markdown('<div class="login-card">', unsafe_allow_html=True)
        
        # Logo
        if LOGO_PATH.exists():
            st.image(str(LOGO_PATH), use_container_width=True)
        else:
            st.markdown(
                "<div style='text-align:center;font-size:3rem;'>🏫</div>",
                unsafe_allow_html=True
            )
        
        st.markdown(
            "<h2>Colegio Pukaray</h2>"
            "<p>Sistema de Convivencia Escolar — Acceso Institucional</p>",
            unsafe_allow_html=True,
        )
        
        # Inicializar error
        if "_login_error" not in st.session_state:
            st.session_state["_login_error"] = False
        
        username = st.text_input("👤 Usuario", key="login_username", placeholder="nombre.usuario")
        password = st.text_input("🔒 Contraseña", type="password", key="login_password")
        
        if st.button("Iniciar Sesión →", use_container_width=True):
            if not username or not password:
                st.warning("Complete todos los campos.")
            else:
                cb_login(username, password)
                if st.session_state.get("_login_error"):
                    st.error("❌ Credenciales incorrectas. Verifique usuario y contraseña.")
                    st.session_state["_login_error"] = False
                else:
                    st.rerun()
        
        st.markdown("---")
        st.markdown(
            "<div style='text-align:center;font-size:0.78rem;color:#888;'>"
            "Acceso restringido a personal autorizado del Colegio Pukaray.<br>"
            "Contraseñas cifradas con SHA-256."
            "</div>",
            unsafe_allow_html=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# MÓDULO 2: FORMULARIOS DE REGISTRO
# ─────────────────────────────────────────────────────────────────────────────

# Opciones de campos de selección
DEPTOS        = ["Orientación", "Inspectoría", "Dirección", "PIE/NEE", "Psicología", "Trabajo Social", "Otro"]
CURSOS        = ["Pre-Kínder", "Kínder", "1°A", "1°B", "2°A", "2°B", "3°A", "3°B",
                 "4°A", "4°B", "5°A", "5°B", "6°A", "6°B", "7°A", "7°B", "8°A", "8°B",
                 "I° Medio A", "I° Medio B", "II° Medio A", "II° Medio B",
                 "III° Medio A", "III° Medio B", "IV° Medio A", "IV° Medio B"]
TIPOS_FALTA   = ["Leve", "Menos Grave", "Grave", "Gravísima", "No Aplica"]
PROTOCOLOS    = ["Ninguno", "Protocolo Acoso Escolar", "Protocolo Violencia Física",
                 "Protocolo Violencia Psicológica", "Protocolo Abuso Sexual",
                 "Protocolo Deserción Escolar", "Protocolo Ausentismo",
                 "Protocolo Sustancias", "Protocolo Familia", "Otro"]
ESTADOS_CASO  = ["Abierto", "En Proceso", "Seguimiento", "Cerrado", "No Aplica"]
CONDICIONES   = ["Inicial", "Activo", "En Revisión", "Resuelto", "Derivado", "Archivado"]


def _safe_bool(val) -> bool:
    """Convierte string 'True'/'False' a bool para precargar checkboxes."""
    return str(val).strip().lower() in ("true", "si", "sí", "1", "yes")


def _safe_index(lista: list, valor: str, default: int = 0) -> int:
    """Retorna el índice de 'valor' en 'lista', o 'default' si no existe."""
    try:
        return lista.index(str(valor))
    except (ValueError, TypeError):
        return default


def pagina_formulario():
    """
    Módulo 2: Formulario de registro de entrevistas.

    MODO NUEVO CC: Crea un registro nuevo en el Excel (append).
    MODO CC EXISTENTE: Al seleccionar un CC, autocompleta todos los campos
        con los datos del último registro de ese CC. El usuario puede
        modificarlos libremente. Al guardar, el comportamiento es:
        - Si el id_registro coincide con uno existente → ACTUALIZA esa fila.
        - Si no → AGREGA una nueva fila (nueva entrevista del mismo CC).
        El nombre del CC no cambia; la entrevista se vincula al mismo caso.

    Manejo de session_state:
    - 'form_preload': dict con datos precargados del CC seleccionado, o None.
    - 'form_cc_prev': CC seleccionado en el ciclo anterior (detecta cambio).
    - Todos los widgets leen su valor desde session_state sólo vía 'value='.
      No se mutan keys de widgets directamente post-instanciación.
    """
    mostrar_header("Registro de Entrevista", "Formulario de Ingreso")

    df = obtener_df()

    # ── Selector de tipo de ficha ──
    tipo_ficha = st.radio(
        "Tipo de Ficha:",
        ["Apoderado / Familia", "Estudiante"],
        horizontal=True,
        key="radio_tipo_ficha"
    )
    tipo_clave = "apoderado" if "Apoderado" in tipo_ficha else "estudiante"

    st.markdown("---")

    # ────────────────────────────────────────────────────────────────────────
    # SECCIÓN CC — va primero para poder precargar el resto del formulario
    # ────────────────────────────────────────────────────────────────────────
    st.markdown('<div class="seccion-card"><h3>🔑 Código de Caso (CC)</h3>', unsafe_allow_html=True)

    cc_modo = st.radio(
        "¿Cómo asignar el CC?",
        ["Generar nuevo CC automáticamente", "Seleccionar CC existente"],
        horizontal=True,
        key="radio_cc_modo"
    )

    cc_final        = ""
    id_registro_editar = None   # ID del registro a actualizar (None = nuevo)
    preload         = {}        # datos que precargarán los campos

    if cc_modo == "Generar nuevo CC automáticamente":
        # Limpiar precarga al volver a modo nuevo
        st.session_state["form_preload"] = None
        st.session_state["form_cc_prev"] = ""
        cc_preview = generar_cc_nuevo(df, "")
        st.info(f"🔄 Se generará: **{cc_preview}**")
        cc_final = cc_preview

    else:
        ccs_existentes = obtener_ccs_existentes(df)
        if not ccs_existentes:
            st.warning("No hay CCs existentes. Se creará uno nuevo automáticamente.")
            st.session_state["form_preload"] = None
            cc_final = generar_cc_nuevo(df, "")
        else:
            cc_selec = st.selectbox(
                "Seleccionar CC existente:", ccs_existentes, key="f_cc_sel"
            )
            cc_final = cc_selec

            # Detectar cambio de CC → recargar preload
            if cc_selec != st.session_state.get("form_cc_prev", ""):
                df_cc_rows = df[df["cc"] == cc_selec]
                if not df_cc_rows.empty:
                    ultimo = df_cc_rows.iloc[-1].to_dict()
                    st.session_state["form_preload"] = ultimo
                else:
                    st.session_state["form_preload"] = None
                st.session_state["form_cc_prev"] = cc_selec

            preload = st.session_state.get("form_preload") or {}

            if preload:
                # Mostrar banner informativo
                n_ents_cc = len(df[df["cc"] == cc_selec])
                st.info(
                    f"📋 CC **{cc_selec}** — Estudiante: **{preload.get('estudiante', '—')}** | "
                    f"{n_ents_cc} entrevista(s) previas. "
                    f"Los campos se completaron con el último registro. Modifíquelos si corresponde."
                )
                id_registro_editar = preload.get("id_registro")

    # Campos de CC adicionales — precargados si hay preload
    col_f1, col_f2, col_n = st.columns(3)
    with col_f1:
        folio = st.text_input(
            "📄 Folio", key="f_folio",
            value=preload.get("folio", "")
        )
    with col_f2:
        # N° entrevista: sugerir el siguiente número al existente
        n_ent_prev = 1
        try:
            n_ent_prev = int(preload.get("n_entrevista", 1)) if preload else 1
        except (ValueError, TypeError):
            n_ent_prev = 1
        n_entrevista = st.number_input(
            "N° Entrevista", min_value=1, step=1, key="f_n_ent",
            value=n_ent_prev
        )
    with col_n:
        idx_estado = _safe_index(ESTADOS_CASO, preload.get("estado_caso", "Abierto"))
        estado_caso = st.selectbox(
            "📊 Estado del Caso", ESTADOS_CASO, index=idx_estado, key="f_estado"
        )

    idx_condicion = _safe_index(CONDICIONES, preload.get("condicion_caso", "Inicial"))
    condicion_caso = st.selectbox(
        "📌 Condición del Caso", CONDICIONES, index=idx_condicion, key="f_condicion"
    )

    st.markdown("</div>", unsafe_allow_html=True)

    # ── SECCIÓN 1: Datos Generales ──
    st.markdown('<div class="seccion-card"><h3>📋 Datos Generales de la Entrevista</h3>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1.2, 1, 1])
    with col1:
        fecha_ent = st.date_input(
            "📅 Fecha de Entrevista", value=date.today(), key="f_fecha"
        )
    with col2:
        hora_ent = st.time_input(
            "🕐 Hora", value=time(9, 0), key="f_hora"
        )
    with col3:
        idx_depto = _safe_index(DEPTOS, preload.get("depto_cita", DEPTOS[0]))
        depto_cita = st.selectbox(
            "🏢 Depto. que cita", DEPTOS, index=idx_depto, key="f_depto"
        )

    col4, col5 = st.columns([1, 2])
    with col4:
        idx_curso = _safe_index(CURSOS, preload.get("curso", CURSOS[0]))
        curso = st.selectbox(
            "📚 Curso", CURSOS, index=idx_curso, key="f_curso"
        )
    with col5:
        estudiante = st.text_input(
            "👤 Nombre completo del Estudiante *", key="f_estudiante",
            value=preload.get("estudiante", "")
        )

    col6, col7 = st.columns([1, 2])
    with col6:
        run_est = st.text_input(
            "🪪 RUN Estudiante", placeholder="12.345.678-9", key="f_run",
            value=preload.get("run_estudiante", "")
        )
    with col7:
        participantes = st.text_input(
            "👥 Participantes de la entrevista",
            placeholder="Ej: Madre, Padre, Orientador, Profesor Jefe",
            key="f_participantes",
            value=preload.get("participantes", "")
        )

    st.markdown("</div>", unsafe_allow_html=True)

    # ── SECCIÓN 3: Clasificación ──
    st.markdown('<div class="seccion-card"><h3>📑 Clasificación y Motivo</h3>', unsafe_allow_html=True)

    motivo = st.text_area(
        "📝 Motivo de la entrevista *", height=90, key="f_motivo",
        value=preload.get("motivo", "")
    )
    col_tf, col_pr = st.columns(2)
    with col_tf:
        idx_tf = _safe_index(TIPOS_FALTA, preload.get("tipo_falta_rice", TIPOS_FALTA[0]))
        tipo_falta = st.selectbox(
            "⚖️ Tipo de Falta RICE", TIPOS_FALTA, index=idx_tf, key="f_tipo_falta"
        )
    with col_pr:
        idx_pr = _safe_index(PROTOCOLOS, preload.get("protocolos", PROTOCOLOS[0]))
        protocolo = st.selectbox(
            "📋 Protocolos asociados", PROTOCOLOS, index=idx_pr, key="f_protocolo"
        )

    st.markdown("</div>", unsafe_allow_html=True)

    # ── SECCIÓN 4: Contenido de la Entrevista ──
    st.markdown('<div class="seccion-card"><h3>📖 Contenido de la Entrevista</h3>', unsafe_allow_html=True)

    antecedentes = st.text_area(
        "🗂️ Antecedentes / Relato de Hechos",
        height=160,
        placeholder="Describa los antecedentes relevantes del caso...",
        key="f_antecedentes",
        value=preload.get("antecedentes", "")
    )
    acuerdos = st.text_area(
        "🤝 Acuerdos o Conclusiones",
        height=130,
        placeholder="Enumere los acuerdos o conclusiones de la entrevista...",
        key="f_acuerdos",
        value=preload.get("acuerdos", "")
    )
    compromisos = st.text_area(
        "✅ Compromisos Adquiridos",
        height=120,
        placeholder="Especifique los compromisos asumidos por cada parte...",
        key="f_compromisos",
        value=preload.get("compromisos", "")
    )
    observaciones = st.text_area(
        "💬 Observaciones Adicionales",
        height=90,
        key="f_observaciones",
        value=preload.get("observaciones", "")
    )
    resumen_libro = st.text_area(
        "📒 Resumen del Libro de Clases",
        height=90,
        placeholder="Resumen de anotaciones relevantes en libro de clases...",
        key="f_resumen_libro",
        value=preload.get("resumen_libro_clases", "")
    )

    st.markdown("</div>", unsafe_allow_html=True)

    # ── SECCIÓN 5: Checklist de Cierre ──
    st.markdown('<div class="seccion-card"><h3>☑️ Checklist de Cierre</h3>', unsafe_allow_html=True)

    col_chk1, col_chk2 = st.columns(2)
    with col_chk1:
        chk_notif    = st.checkbox("📬 Notificación realizada",     key="chk_notif",
                                   value=_safe_bool(preload.get("check_notificacion", False)))
        chk_firma_ap = st.checkbox("✍️ Firma apoderado",            key="chk_firma_ap",
                                   value=_safe_bool(preload.get("check_firma_apoderado", False)))
        chk_firma_st = st.checkbox("✍️ Firma estudiante",           key="chk_firma_st",
                                   value=_safe_bool(preload.get("check_firma_estudiante", False)))
        chk_firma_dc = st.checkbox("✍️ Firma docente/encargado",    key="chk_firma_dc",
                                   value=_safe_bool(preload.get("check_firma_docente", False)))
    with col_chk2:
        chk_deriv  = st.checkbox("➡️ Derivación realizada",         key="chk_deriv",
                                 value=_safe_bool(preload.get("check_derivacion", False)))
        chk_seguim = st.checkbox("🔄 Seguimiento programado",       key="chk_seguim",
                                 value=_safe_bool(preload.get("check_seguimiento", False)))
        chk_cierre = st.checkbox("🔒 Cierre de CC",                 key="chk_cierre",
                                 value=_safe_bool(preload.get("check_cierre_cc", False)))
        chk_copia  = st.checkbox("📄 Copia entregada a apoderado",  key="chk_copia",
                                 value=_safe_bool(preload.get("check_copia_apoderado", False)))

    st.markdown("</div>", unsafe_allow_html=True)

    # ── ACCIONES ──
    st.markdown("---")

    if not estudiante.strip():
        st.warning("⚠️ El campo **Nombre del Estudiante** es obligatorio.")

    col_btn1, col_btn2, col_esp = st.columns([1.5, 1.5, 2])

    with col_btn1:
        guardar_clicked = st.button(
            "💾 Guardar Registro",
            use_container_width=True,
            disabled=not estudiante.strip()
        )

    with col_btn2:
        generar_word_clicked = st.button(
            "📄 Generar Acta Word",
            use_container_width=True,
            disabled=not estudiante.strip()
        )

    # ── Procesamiento: Guardar ──
    if guardar_clicked and estudiante.strip():
        # Construir el dict del registro con todos los campos
        reg_data = {
            "tipo_ficha":            tipo_clave,
            "fecha":                 str(fecha_ent),
            "hora":                  str(hora_ent),
            "depto_cita":            depto_cita,
            "curso":                 curso,
            "estudiante":            estudiante.strip(),
            "run_estudiante":        run_est.strip(),
            "participantes":         participantes.strip(),
            "motivo":                motivo.strip(),
            "tipo_falta_rice":       tipo_falta,
            "protocolos":            protocolo,
            "cc":                    cc_final,
            "folio":                 folio.strip(),
            "n_entrevista":          str(int(n_entrevista)),
            "estado_caso":           estado_caso,
            "condicion_caso":        condicion_caso,
            "antecedentes":          antecedentes.strip(),
            "acuerdos":              acuerdos.strip(),
            "compromisos":           compromisos.strip(),
            "observaciones":         observaciones.strip(),
            "resumen_libro_clases":  resumen_libro.strip(),
            "check_notificacion":    str(chk_notif),
            "check_firma_apoderado": str(chk_firma_ap),
            "check_firma_estudiante":str(chk_firma_st),
            "check_firma_docente":   str(chk_firma_dc),
            "check_derivacion":      str(chk_deriv),
            "check_seguimiento":     str(chk_seguim),
            "check_cierre_cc":       str(chk_cierre),
            "check_copia_apoderado": str(chk_copia),
            "usuario_registro":      st.session_state["usuario"]["username"],
            "timestamp_registro":    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }

        # ── Lógica de UPDATE vs INSERT ──────────────────────────────────────
        # Si venimos de CC existente y el id_registro_editar está en el df,
        # actualizamos esa fila en lugar de añadir una nueva.
        # Esto evita duplicados cuando el usuario solo edita una entrevista.
        existe_id = (
            id_registro_editar is not None
            and not df.empty
            and id_registro_editar in df["id_registro"].values
        )

        if existe_id:
            # ACTUALIZAR fila existente
            reg_data["id_registro"] = id_registro_editar
            idx_fila = df.index[df["id_registro"] == id_registro_editar][0]
            for col, val in reg_data.items():
                df.at[idx_fila, col] = val
            df_nuevo     = df.copy()
            accion_label = f"✅ Registro **{id_registro_editar}** actualizado. CC: **{cc_final}**"
        else:
            # INSERTAR nuevo registro
            nuevo_id          = generar_id_registro(df)
            reg_data["id_registro"] = nuevo_id
            df_nuevo          = pd.concat(
                [df, pd.DataFrame([reg_data])], ignore_index=True
            )
            accion_label = f"✅ Nuevo registro **{nuevo_id}** guardado. CC: **{cc_final}**"

        if guardar_registro(df_nuevo):
            recargar_df()
            # Actualizar preload con los datos recién guardados
            st.session_state["form_preload"]    = reg_data
            st.session_state["form_guardado"]   = True
            st.success(accion_label)

    # ── Procesamiento: Generar Word ──
    if generar_word_clicked and estudiante.strip():
        datos_word = {
            "id_registro":            id_registro_editar or generar_id_registro(df),
            "tipo_ficha":             tipo_clave,
            "fecha":                  str(fecha_ent),
            "hora":                   str(hora_ent),
            "depto_cita":             depto_cita,
            "curso":                  curso,
            "estudiante":             estudiante.strip(),
            "run_estudiante":         run_est.strip(),
            "participantes":          participantes.strip(),
            "motivo":                 motivo.strip(),
            "tipo_falta_rice":        tipo_falta,
            "protocolos":             protocolo,
            "cc":                     cc_final,
            "folio":                  folio.strip(),
            "n_entrevista":           str(int(n_entrevista)),
            "estado_caso":            estado_caso,
            "condicion_caso":         condicion_caso,
            "antecedentes":           antecedentes.strip(),
            "acuerdos":               acuerdos.strip(),
            "compromisos":            compromisos.strip(),
            "observaciones":          observaciones.strip(),
            "resumen_libro_clases":   resumen_libro.strip(),
            "check_notificacion":     str(chk_notif),
            "check_firma_apoderado":  str(chk_firma_ap),
            "check_firma_estudiante": str(chk_firma_st),
            "check_firma_docente":    str(chk_firma_dc),
            "check_derivacion":       str(chk_deriv),
            "check_seguimiento":      str(chk_seguim),
            "check_cierre_cc":        str(chk_cierre),
            "check_copia_apoderado":  str(chk_copia),
        }

        word_bytes = generar_word(datos_word, tipo_clave)

        if word_bytes:
            nombre_archivo = (
                f"Acta_{tipo_clave.upper()}_{cc_final}_{fecha_ent.strftime('%Y%m%d')}.docx"
            )
            st.download_button(
                label="⬇️ Descargar Acta Word",
                data=word_bytes,
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )


# ─────────────────────────────────────────────────────────────────────────────
# MÓDULO 3: GESTIÓN DE CASOS POR CC
# ─────────────────────────────────────────────────────────────────────────────

def pagina_gestion_cc():
    """
    Módulo 3: Visualización de casos por Código CC.
    Muestra la línea de tiempo de entrevistas asociadas al CC seleccionado
    y el estado actual del caso.
    """
    mostrar_header("Gestión de Casos (CC)", "Línea de Tiempo")
    
    df = obtener_df()
    
    if df.empty:
        st.info("📭 No hay entrevistas registradas aún.")
        return
    
    ccs = obtener_ccs_existentes(df)
    if not ccs:
        st.info("No hay Códigos de Caso registrados.")
        return
    
    col_sel, col_info = st.columns([1, 2])
    
    with col_sel:
        st.markdown('<div class="seccion-card"><h3>🔍 Seleccionar Caso</h3>', unsafe_allow_html=True)
        cc_sel = st.selectbox("Código de Caso (CC):", ccs, key="cc_gestion_sel")
        
        df_cc  = df[df["cc"] == cc_sel].copy()
        
        # Estado del caso (el más reciente)
        if not df_cc.empty:
            ult_estado = df_cc["estado_caso"].iloc[-1] if "estado_caso" in df_cc.columns else "—"
            ult_cond   = df_cc["condicion_caso"].iloc[-1] if "condicion_caso" in df_cc.columns else "—"
            ult_est    = df_cc["estudiante"].iloc[-1] if "estudiante" in df_cc.columns else "—"
            ult_curso  = df_cc["curso"].iloc[-1] if "curso" in df_cc.columns else "—"
            
            st.markdown("**Estudiante:**")
            st.markdown(f"👤 {ult_est} ({ult_curso})")
            
            st.markdown("**Estado actual:**")
            st.markdown(badge_estado(ult_estado), unsafe_allow_html=True)
            
            st.markdown("**Condición:**")
            st.info(ult_cond)
            
            st.markdown(f"**Total entrevistas:** {len(df_cc)}")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    with col_info:
        st.markdown('<div class="seccion-card"><h3>📅 Línea de Tiempo del Caso</h3>', unsafe_allow_html=True)
        
        if df_cc.empty:
            st.warning("Sin entrevistas para este CC.")
        else:
            # Ordenar por fecha
            try:
                df_cc["fecha_dt"] = pd.to_datetime(df_cc["fecha"], errors="coerce")
                df_cc = df_cc.sort_values("fecha_dt")
            except Exception:
                pass
            
            for _, row in df_cc.iterrows():
                tipo_icon = "👨‍👩‍👧" if str(row.get("tipo_ficha", "")).lower() == "apoderado" else "🎒"
                est_html  = badge_estado(row.get("estado_caso", ""))
                
                st.markdown(
                    f"""
                    <div class="timeline-item">
                        <div class="timeline-fecha">
                            {tipo_icon}<br>
                            {row.get("fecha", "")}
                        </div>
                        <div>
                            <strong>ID:</strong> {row.get("id_registro", "—")} |
                            <strong>N°:</strong> {row.get("n_entrevista", "—")}<br>
                            <strong>Motivo:</strong> {str(row.get("motivo", ""))[:80]}...<br>
                            <strong>Estado:</strong> {est_html} &nbsp;
                            <strong>Folio:</strong> {row.get("folio", "—")}<br>
                            <em style="font-size:0.8rem;color:#666;">
                                Registrado por: {row.get("usuario_registro", "—")} —
                                {row.get("timestamp_registro", "—")}
                            </em>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # ── Estadísticas del CC ──
    if not df_cc.empty:
        st.markdown("---")
        st.markdown('<div class="seccion-card"><h3>📊 Resumen del Caso</h3>', unsafe_allow_html=True)
        
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.metric("Total Entrevistas", len(df_cc))
        with c2:
            tipos = df_cc["tipo_ficha"].value_counts().to_dict()
            st.metric("Entrevistas Apoderado", tipos.get("apoderado", 0))
        with c3:
            st.metric("Entrevistas Estudiante", tipos.get("estudiante", 0))
        with c4:
            primera = df_cc["fecha"].min() if not df_cc.empty else "—"
            st.metric("Fecha Apertura", str(primera))
        
        st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# MÓDULO 4: REVISIÓN DE ENTREVISTAS
# ─────────────────────────────────────────────────────────────────────────────

def pagina_revision():
    """
    Módulo 4: Revisión y consulta de entrevistas registradas.

    Tabla de selección compacta: sólo CC | Estudiante | Curso | Tipo Registro.
    Al marcar una fila se despliega la ficha completa en modo solo lectura.
    Las etiquetas de visualización usan nombres legibles (sin guiones bajos).
    Los campos de texto largo (antecedentes, acuerdos, compromisos) siempre
    muestran el valor real leído del DataFrame, no del session_state de otro widget.
    """
    mostrar_header("Revisión de Entrevistas", "Modo Solo Lectura")

    df = obtener_df()

    if df.empty:
        st.info("📭 No hay entrevistas registradas aún.")
        return

    # ── Filtros ──
    st.markdown('<div class="seccion-card"><h3>🔍 Filtros de Búsqueda</h3>', unsafe_allow_html=True)

    col_f1, col_f2, col_f3, col_f4 = st.columns(4)
    with col_f1:
        filtro_tipo = st.selectbox(
            "Tipo de Registro:", ["Todos", "apoderado", "estudiante"], key="rev_tipo"
        )
    with col_f2:
        filtro_estado = st.selectbox(
            "Estado del Caso:", ["Todos"] + ESTADOS_CASO, key="rev_estado"
        )
    with col_f3:
        filtro_curso = st.selectbox(
            "Curso:", ["Todos"] + CURSOS, key="rev_curso"
        )
    with col_f4:
        filtro_texto = st.text_input("Buscar por estudiante o CC:", key="rev_texto")

    st.markdown("</div>", unsafe_allow_html=True)

    # Aplicar filtros
    df_filtrado = df.copy()

    if filtro_tipo != "Todos":
        df_filtrado = df_filtrado[
            df_filtrado["tipo_ficha"].str.lower() == filtro_tipo.lower()
        ]
    if filtro_estado != "Todos":
        df_filtrado = df_filtrado[
            df_filtrado["estado_caso"].str.lower() == filtro_estado.lower()
        ]
    if filtro_curso != "Todos":
        df_filtrado = df_filtrado[df_filtrado["curso"] == filtro_curso]
    if filtro_texto.strip():
        mask = (
            df_filtrado["estudiante"].str.contains(filtro_texto, case=False, na=False)
            | df_filtrado["cc"].str.contains(filtro_texto, case=False, na=False)
        )
        df_filtrado = df_filtrado[mask]

    n_encontrados = len(df_filtrado)
    st.markdown(f"**{n_encontrados}** entrevista(s) encontrada(s).")

    if df_filtrado.empty:
        st.warning("No hay registros que coincidan con los filtros.")
        return

    # ────────────────────────────────────────────────────────────────────────
    # TABLA DE SELECCIÓN COMPACTA
    # Sólo 4 columnas: CC | Estudiante | Curso | Tipo Registro
    # El id_registro se conserva como columna oculta para la recuperación
    # del registro completo, pero NO se muestra en la tabla.
    # ────────────────────────────────────────────────────────────────────────
    st.markdown(
        '<div class="seccion-card"><h3>📋 Seleccionar Entrevista</h3>',
        unsafe_allow_html=True
    )
    st.caption("Marque la casilla de la entrevista que desea revisar.")

    # Construir tabla compacta con las 4 columnas visibles + id oculto
    df_compacto = df_filtrado[
        [c for c in ["id_registro", "cc", "estudiante", "curso", "tipo_ficha"]
         if c in df_filtrado.columns]
    ].copy().reset_index(drop=True)

    # Renombrar para visualización legible
    df_compacto.rename(columns={"tipo_ficha": "Tipo Registro"}, inplace=True)

    # Insertar columna de selección al inicio
    df_compacto.insert(0, "Ver", False)

    df_editado = st.data_editor(
        df_compacto,
        column_config={
            "Ver":          st.column_config.CheckboxColumn("✓", width="small"),
            "cc":           st.column_config.TextColumn("CC"),
            "estudiante":   st.column_config.TextColumn("Estudiante"),
            "curso":        st.column_config.TextColumn("Curso"),
            "Tipo Registro": st.column_config.TextColumn("Tipo Registro"),
            # Ocultar id_registro de la vista pero mantenerlo en el df
            "id_registro":  st.column_config.TextColumn("ID", disabled=True),
        },
        column_order=["Ver", "cc", "estudiante", "curso", "Tipo Registro"],
        use_container_width=True,
        hide_index=True,
        key="tabla_revision",
        disabled=["cc", "estudiante", "curso", "Tipo Registro", "id_registro"],
    )

    st.markdown("</div>", unsafe_allow_html=True)

    # ── Recuperar el registro seleccionado ──
    seleccionados = df_editado[df_editado["Ver"] == True]

    if seleccionados.empty:
        st.info("ℹ️ Marque una fila de la tabla para ver la ficha completa.")
        return

    if len(seleccionados) > 1:
        st.warning("⚠️ Seleccione solo una entrevista para ver el detalle.")
        return

    id_sel = seleccionados.iloc[0]["id_registro"]
    # Leer siempre desde df_filtrado (fuente de verdad), NO desde el editor
    filas = df_filtrado[df_filtrado["id_registro"] == id_sel]
    if filas.empty:
        st.error("No se encontró el registro seleccionado.")
        return
    reg = filas.iloc[0]

    # ────────────────────────────────────────────────────────────────────────
    # FICHA COMPLETA EN MODO SOLO LECTURA
    # Etiquetas legibles (sin guiones bajos, sin códigos internos)
    # ────────────────────────────────────────────────────────────────────────
    st.markdown("---")
    tipo_label = "Apoderado / Familia" if str(reg.get("tipo_ficha", "")).lower() == "apoderado" else "Estudiante"
    st.markdown(
        f'<div class="seccion-card">'
        f'<h3>📄 Ficha Completa — {id_sel} &nbsp;|&nbsp; {tipo_label} &nbsp;(Solo Lectura)</h3>',
        unsafe_allow_html=True
    )

    # ── Bloque 1: Identificación del caso ──
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"**Código de Caso (CC):** `{reg.get('cc', '—')}`")
        st.markdown(f"**Folio:** {reg.get('folio', '—')}")
        st.markdown(f"**N° de Entrevista:** {reg.get('n_entrevista', '—')}")
    with c2:
        st.markdown(f"**Fecha de Entrevista:** {reg.get('fecha', '—')}")
        st.markdown(f"**Hora:** {reg.get('hora', '—')}")
        st.markdown(f"**Departamento que cita:** {reg.get('depto_cita', '—')}")
    with c3:
        st.markdown("**Estado del Caso:**")
        st.markdown(badge_estado(reg.get("estado_caso", "")), unsafe_allow_html=True)
        st.markdown(f"**Condición del Caso:** {reg.get('condicion_caso', '—')}")

    st.markdown("---")

    # ── Bloque 2: Datos del estudiante ──
    c4, c5 = st.columns(2)
    with c4:
        st.markdown(f"**Estudiante:** {reg.get('estudiante', '—')}")
        st.markdown(f"**RUN del Estudiante:** {reg.get('run_estudiante', '—')}")
        st.markdown(f"**Curso:** {reg.get('curso', '—')}")
    with c5:
        st.markdown("**Participantes de la Entrevista:**")
        st.markdown(reg.get("participantes", "—") or "—")

    st.markdown("**Motivo de la Entrevista:**")
    st.markdown(f"> {reg.get('motivo', '—') or '—'}")
    st.markdown(
        f"**Tipo de Falta RICE:** {reg.get('tipo_falta_rice', '—')} &nbsp;|&nbsp; "
        f"**Protocolo Asociado:** {reg.get('protocolos', '—')}",
        unsafe_allow_html=True
    )

    st.markdown("---")

    # ── Bloque 3: Contenido de la entrevista ──
    # IMPORTANTE: los text_area usan unique keys que incluyen el id_sel
    # para evitar conflictos de session_state entre distintas selecciones.
    # El valor se lee directamente de `reg` (Serie del DataFrame), no de widgets
    # del formulario de registro, garantizando que antecedentes/acuerdos
    # siempre reflejen los datos guardados en disco.

    st.markdown("**📑 Antecedentes / Relato de Hechos:**")
    val_ant = str(reg.get("antecedentes", "") or "")
    st.text_area(
        "Antecedentes", value=val_ant if val_ant else "(Sin registro)",
        height=130, disabled=True,
        key=f"rv_ant_{id_sel}", label_visibility="collapsed"
    )

    st.markdown("**🤝 Acuerdos o Conclusiones:**")
    val_acu = str(reg.get("acuerdos", "") or "")
    st.text_area(
        "Acuerdos", value=val_acu if val_acu else "(Sin registro)",
        height=110, disabled=True,
        key=f"rv_acu_{id_sel}", label_visibility="collapsed"
    )

    st.markdown("**✅ Compromisos Adquiridos:**")
    val_comp = str(reg.get("compromisos", "") or "")
    st.text_area(
        "Compromisos", value=val_comp if val_comp else "(Sin registro)",
        height=110, disabled=True,
        key=f"rv_comp_{id_sel}", label_visibility="collapsed"
    )

    st.markdown("**💬 Observaciones Adicionales:**")
    val_obs = str(reg.get("observaciones", "") or "")
    st.text_area(
        "Observaciones", value=val_obs if val_obs else "(Sin registro)",
        height=80, disabled=True,
        key=f"rv_obs_{id_sel}", label_visibility="collapsed"
    )

    st.markdown("**📒 Resumen del Libro de Clases:**")
    val_lib = str(reg.get("resumen_libro_clases", "") or "")
    st.text_area(
        "Resumen", value=val_lib if val_lib else "(Sin registro)",
        height=80, disabled=True,
        key=f"rv_lib_{id_sel}", label_visibility="collapsed"
    )

    st.markdown("---")

    # ── Bloque 4: Checklist de cierre ──
    st.markdown("**☑️ Checklist de Cierre:**")

    checks = {
        "Notificación realizada":       reg.get("check_notificacion", "False"),
        "Firma apoderado":              reg.get("check_firma_apoderado", "False"),
        "Firma estudiante":             reg.get("check_firma_estudiante", "False"),
        "Firma docente / encargado":    reg.get("check_firma_docente", "False"),
        "Derivación realizada":         reg.get("check_derivacion", "False"),
        "Seguimiento programado":       reg.get("check_seguimiento", "False"),
        "Cierre de CC":                 reg.get("check_cierre_cc", "False"),
        "Copia entregada a apoderado":  reg.get("check_copia_apoderado", "False"),
    }

    cols_chk = st.columns(4)
    for i, (item, val) in enumerate(checks.items()):
        ok = str(val).strip().lower() in ("true", "si", "sí", "1", "yes")
        simbolo = "✅" if ok else "⬜"
        with cols_chk[i % 4]:
            st.markdown(f"{simbolo} {item}")

    st.markdown(
        f"<br><em style='font-size:0.8rem;color:#888;'>"
        f"Registrado por: {reg.get('usuario_registro', '—')} &nbsp;—&nbsp; "
        f"Fecha de registro en sistema: {reg.get('timestamp_registro', '—')}"
        f"</em>",
        unsafe_allow_html=True
    )

    st.markdown("</div>", unsafe_allow_html=True)

    # ── Descarga del acta ──
    if st.button("📄 Descargar Acta de esta Entrevista", key=f"btn_dl_rev_{id_sel}"):
        datos_word = {k: str(v) for k, v in reg.items()}
        word_bytes = generar_word(datos_word, str(reg.get("tipo_ficha", "apoderado")))
        if word_bytes:
            st.download_button(
                "⬇️ Descargar Acta Word",
                data=word_bytes,
                file_name=f"Acta_{id_sel}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_word_{id_sel}"
            )


# ─────────────────────────────────────────────────────────────────────────────
# MÓDULO 5: BORRADO SEGURO DE REGISTROS
# ─────────────────────────────────────────────────────────────────────────────

def pagina_borrado_seguro():
    """
    Módulo 5: Borrado seguro de entrevistas (NO de casos completos CC).

    Reglas de seguridad implementadas:
    1. El usuario filtra por CC o estudiante y elige registros a eliminar
       mediante un multiselect compacto (ID + resumen de 1 línea).
       Se elimina la tabla completa redundante: los datos completos ya se
       pueden revisar en el módulo de Revisión de Entrevistas.
    2. El usuario debe escribir explícitamente 'BORRAR' para confirmar.
    3. RESPALDO AUTOMÁTICO antes de cualquier eliminación (regla inquebrantable).
    4. Solo se eliminan las entrevistas seleccionadas, NO el CC completo.

    Solo accesible para usuarios con rol 'administrador'.
    """
    mostrar_header("Borrado Seguro de Registros", "⚠️ Acción Irreversible")

    usuario = st.session_state.get("usuario", {})
    if usuario.get("rol") != "administrador":
        st.error("🚫 Acceso denegado. Solo el Administrador puede eliminar registros.")
        return

    st.markdown(
        """
        <div style="background:#FFF3CD;border:1px solid #856404;border-left:5px solid #856404;
                    border-radius:4px;padding:1rem;margin-bottom:1rem;">
        ⚠️ <strong>ZONA DE ALTA RESTRICCIÓN</strong><br>
        Esta acción eliminará entrevistas individuales del registro.
        <strong>NO</strong> elimina el Código de Caso (CC) completo.
        Un respaldo automático se crea antes de cada operación.
        Para revisar el detalle completo de una entrevista antes de borrarla,
        use el módulo <em>Revisión de Entrevistas</em>.
        </div>
        """,
        unsafe_allow_html=True,
    )

    df = obtener_df()

    if df.empty:
        st.info("No hay registros para eliminar.")
        return

    # ────────────────────────────────────────────────────────────────────────
    # SELECTOR COMPACTO — sin tabla completa redundante
    # El usuario filtra primero, luego elige por ID con descripción resumida.
    # ────────────────────────────────────────────────────────────────────────
    st.markdown(
        '<div class="seccion-card"><h3>🗑️ Seleccionar Entrevistas a Eliminar</h3>',
        unsafe_allow_html=True
    )

    # Filtro previo para reducir la lista (especialmente útil con muchos registros)
    col_fb1, col_fb2 = st.columns(2)
    with col_fb1:
        filtro_cc_b = st.text_input(
            "Filtrar por CC:", key="bor_filtro_cc", placeholder="Ej: CC-7A-2024-001"
        )
    with col_fb2:
        filtro_est_b = st.text_input(
            "Filtrar por estudiante:", key="bor_filtro_est", placeholder="Nombre parcial"
        )

    df_bor = df.copy()
    if filtro_cc_b.strip():
        df_bor = df_bor[df_bor["cc"].str.contains(filtro_cc_b.strip(), case=False, na=False)]
    if filtro_est_b.strip():
        df_bor = df_bor[df_bor["estudiante"].str.contains(filtro_est_b.strip(), case=False, na=False)]

    if df_bor.empty:
        st.warning("Ningún registro coincide con los filtros.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    # Construir etiquetas descriptivas para el multiselect
    # Formato: ENT-001 | CC-7A-2024-001 | Juan Pérez | apoderado | 2024-03-15
    def _etiqueta(row) -> str:
        return (
            f"{row.get('id_registro','?')}  ·  "
            f"{row.get('cc','—')}  ·  "
            f"{row.get('estudiante','—')}  ·  "
            f"{row.get('tipo_ficha','—')}  ·  "
            f"{row.get('fecha','—')}"
        )

    opciones_etiquetas = [_etiqueta(r) for _, r in df_bor.iterrows()]
    # Mapa etiqueta → id_registro
    mapa_etiqueta_id   = {_etiqueta(r): r["id_registro"] for _, r in df_bor.iterrows()}

    seleccion_etiquetas = st.multiselect(
        f"Elija una o más entrevistas ({len(df_bor)} disponibles con los filtros):",
        options=opciones_etiquetas,
        key="bor_multiselect",
        placeholder="Seleccione entrevistas para eliminar…"
    )

    st.markdown("</div>", unsafe_allow_html=True)

    ids_a_borrar = [mapa_etiqueta_id[e] for e in seleccion_etiquetas if e in mapa_etiqueta_id]

    if not ids_a_borrar:
        st.info("ℹ️ Seleccione al menos una entrevista de la lista.")
        return

    # Resumen de lo seleccionado
    st.warning(
        f"⚠️ Ha seleccionado **{len(ids_a_borrar)}** entrevista(s) para eliminar: "
        + ", ".join([f"`{i}`" for i in ids_a_borrar])
    )

    # ── Confirmación explícita ──
    st.markdown(
        '<div class="seccion-card"><h3>🔐 Confirmación de Seguridad</h3>',
        unsafe_allow_html=True
    )
    st.markdown(
        "Para confirmar la eliminación, escriba exactamente la palabra **BORRAR** "
        "en el campo a continuación:"
    )

    confirmacion = st.text_input(
        "Confirmación:", key="confirm_borrar", placeholder="BORRAR"
    )
    st.markdown("</div>", unsafe_allow_html=True)

    col_btn, col_esp = st.columns([1, 3])
    with col_btn:
        ejecutar = st.button(
            "🗑️ EJECUTAR BORRADO",
            use_container_width=True,
            disabled=(confirmacion != "BORRAR"),
            key="btn_ejecutar_borrado"
        )

    if ejecutar and confirmacion == "BORRAR":
        # RESPALDO AUTOMÁTICO (regla inquebrantable)
        ruta_bkp = crear_respaldo_excel()
        if ruta_bkp:
            st.markdown(
                f'<div class="alerta-respaldo">🔒 Respaldo creado ANTES del borrado: '
                f'<code>{Path(ruta_bkp).name}</code></div>',
                unsafe_allow_html=True
            )

        df_nuevo = df[~df["id_registro"].isin(ids_a_borrar)].copy().reset_index(drop=True)

        if guardar_registro(df_nuevo):
            recargar_df()
            st.success(
                f"✅ Se eliminaron {len(ids_a_borrar)} entrevista(s) correctamente. "
                f"Quedan {len(df_nuevo)} registros."
            )
            st.markdown(
                f"**Log de eliminación:** "
                f"{datetime.now().strftime('%d/%m/%Y %H:%M:%S')} | "
                f"Usuario: {usuario['username']} | "
                f"IDs eliminados: {', '.join(ids_a_borrar)}"
            )


# ─────────────────────────────────────────────────────────────────────────────
# MÓDULO 6: RECURSOS NORMATIVOS
# ─────────────────────────────────────────────────────────────────────────────

def pagina_recursos():
    """
    Módulo 6: Acceso a recursos normativos institucionales.
    Enlace externo a NotebookLM normativo.
    """
    mostrar_header("Recursos Normativos", "Marco Legal y Reglamentario")
    
    st.markdown('<div class="seccion-card"><h3>📚 Base Normativa Institucional</h3>', unsafe_allow_html=True)
    
    st.markdown(
        """
        Acceda a la base normativa completa del Colegio Pukaray en materia de
        Convivencia Escolar, incluyendo el Reglamento Interno, protocolos RICE,
        Ley N° 21.128, Ley de Inclusión y documentos complementarios.
        """
    )
    
    st.markdown("---")
    
    # ── Enlace principal NotebookLM ──
    st.markdown(
        """
        <div style="text-align:center;padding:2rem;">
            <a href="https://notebooklm.google.com/notebook/95aeeaf6-1f2a-4b4e-b7b1-b9d2ab3e10d9"
               target="_blank"
               style="
                   display:inline-block;
                   background:#1B4332;
                   color:#FAF3E0;
                   padding:14px 32px;
                   border-radius:8px;
                   font-size:1.1rem;
                   font-weight:700;
                   text-decoration:none;
                   border:2px solid #6B2737;
                   font-family:'Georgia',serif;
                   transition:background 0.2s;
               ">
                📚 Abrir NotebookLM Normativo
            </a>
            <p style="margin-top:1rem;color:#666;font-size:0.88rem;">
                Se abrirá en una nueva pestaña del navegador.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # ── Referencias normativas ──
    st.markdown('<div class="seccion-card"><h3>📋 Referencias Normativas Clave</h3>', unsafe_allow_html=True)
    
    normas = [
        ("Ley N° 20.536", "Sobre Violencia Escolar (2011)"),
        ("Ley N° 21.128", "Aula Segura (2019)"),
        ("Decreto N° 240/1999", "Reglamento de Convivencia"),
        ("Decreto N° 524/1990", "Reglamento de Centros de Alumnos"),
        ("Ord. N° 768/2013", "Orientaciones sobre convivencia escolar"),
        ("Ley N° 20.845", "Inclusión Escolar (2015)"),
        ("Ley N° 21.430", "Protección de la infancia"),
        ("Circular N° 1/2022", "Orientaciones de convivencia escolar MINEDUC"),
    ]
    
    for codigo, descripcion in normas:
        st.markdown(f"- **{codigo}** — {descripcion}")
    
    st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# MÓDULO ADMINISTRACIÓN
# ─────────────────────────────────────────────────────────────────────────────

def pagina_administracion():
    """
    Módulo de administración: solo para rol 'administrador'.
    - Gestión de usuarios (crear, ver, eliminar).
    - Listado de respaldos.
    - Estadísticas generales.
    """
    mostrar_header("Administración del Sistema", "Panel de Control")
    
    usuario = st.session_state.get("usuario", {})
    if usuario.get("rol") != "administrador":
        st.error("🚫 Acceso denegado.")
        return
    
    tab1, tab2, tab3 = st.tabs(["👥 Usuarios", "🔒 Respaldos", "📊 Estadísticas"])
    
    # ── Tab 1: Gestión de usuarios ──
    with tab1:
        st.markdown('<div class="seccion-card"><h3>👥 Usuarios del Sistema</h3>', unsafe_allow_html=True)
        
        df_users = obtener_todos_los_usuarios()
        st.dataframe(df_users, use_container_width=True, hide_index=True)
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        st.markdown('<div class="seccion-card"><h3>➕ Crear Nuevo Usuario</h3>', unsafe_allow_html=True)
        
        cu1, cu2 = st.columns(2)
        with cu1:
            nuevo_username = st.text_input("Nombre de usuario:", key="admin_new_user")
            nuevo_nombre   = st.text_input("Nombre completo:", key="admin_new_nombre")
        with cu2:
            nuevo_pass     = st.text_input("Contraseña:", type="password", key="admin_new_pass")
            nuevo_rol      = st.selectbox("Rol:", ["usuario", "administrador"], key="admin_new_rol")
        
        if st.button("Crear Usuario", key="btn_crear_usuario"):
            if nuevo_username and nuevo_pass and nuevo_nombre:
                ok, msg = crear_usuario(nuevo_username, nuevo_pass, nuevo_rol, nuevo_nombre)
                if ok:
                    st.success(f"✅ {msg}")
                    st.rerun()
                else:
                    st.error(f"❌ {msg}")
            else:
                st.warning("Complete todos los campos.")
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Eliminar usuario
        st.markdown('<div class="seccion-card"><h3>🗑️ Eliminar Usuario</h3>', unsafe_allow_html=True)
        
        if not df_users.empty:
            opciones = {
                f"{r['username']} ({r['nombre']})": r["id"]
                for _, r in df_users.iterrows()
                if r["username"] != usuario["username"]  # no puede eliminarse a sí mismo
            }
            if opciones:
                sel_del = st.selectbox("Seleccionar usuario a eliminar:", list(opciones.keys()), key="admin_del_user")
                if st.button("Eliminar Usuario", key="btn_del_usuario"):
                    if eliminar_usuario(opciones[sel_del]):
                        st.success("✅ Usuario eliminado.")
                        st.rerun()
                    else:
                        st.error("❌ No se pudo eliminar el usuario.")
            else:
                st.info("No hay otros usuarios para eliminar.")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # ── Tab 2: Respaldos ──
    with tab2:
        st.markdown('<div class="seccion-card"><h3>🔒 Archivos de Respaldo</h3>', unsafe_allow_html=True)
        
        respaldos = sorted(RESPALDOS_DIR.glob("*.xlsx"), reverse=True)
        
        if not respaldos:
            st.info("No hay respaldos generados aún.")
        else:
            st.markdown(f"**Total de respaldos:** {len(respaldos)}")
            for r in respaldos:
                stat   = r.stat()
                tamaño = f"{stat.st_size / 1024:.1f} KB"
                fecha_m = datetime.fromtimestamp(stat.st_mtime).strftime("%d/%m/%Y %H:%M:%S")
                st.markdown(
                    f"- 📁 `{r.name}` — {tamaño} — {fecha_m}"
                )
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # ── Tab 3: Estadísticas ──
    with tab3:
        df = obtener_df()
        
        st.markdown('<div class="seccion-card"><h3>📊 Estadísticas Generales</h3>', unsafe_allow_html=True)
        
        if df.empty:
            st.info("Sin datos para mostrar estadísticas.")
        else:
            c1, c2, c3, c4 = st.columns(4)
            with c1:
                st.metric("Total Entrevistas", len(df))
            with c2:
                n_cc = df["cc"].nunique() if "cc" in df.columns else 0
                st.metric("Códigos CC únicos", n_cc)
            with c3:
                n_est = df["estudiante"].nunique() if "estudiante" in df.columns else 0
                st.metric("Estudiantes únicos", n_est)
            with c4:
                if "estado_caso" in df.columns:
                    n_abiertos = len(df[df["estado_caso"].str.lower() == "abierto"])
                    st.metric("Casos Abiertos", n_abiertos)
            
            # Distribución por estado
            if "estado_caso" in df.columns and not df["estado_caso"].dropna().empty:
                st.markdown("**Distribución por Estado:**")
                dist_estado = df["estado_caso"].value_counts().reset_index()
                dist_estado.columns = ["Estado", "Cantidad"]
                st.dataframe(dist_estado, use_container_width=True, hide_index=True)
            
            # Distribución por tipo
            if "tipo_ficha" in df.columns:
                dist_tipo = df["tipo_ficha"].value_counts().reset_index()
                dist_tipo.columns = ["Tipo", "Cantidad"]
                st.markdown("**Por Tipo de Ficha:**")
                st.dataframe(dist_tipo, use_container_width=True, hide_index=True)
        
        st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR DE NAVEGACIÓN
# ─────────────────────────────────────────────────────────────────────────────

def renderizar_sidebar():
    """
    Renderiza el sidebar de navegación.
    Solo visible cuando el usuario está autenticado.
    """
    if not st.session_state.get("autenticado"):
        return
    
    usuario = st.session_state.get("usuario", {})
    
    with st.sidebar:
        # Logo en sidebar
        if LOGO_PATH.exists():
            st.image(str(LOGO_PATH), use_container_width=True)
        else:
            st.markdown(
                "<div style='text-align:center;font-size:2.5rem;'>🏫</div>",
                unsafe_allow_html=True
            )
        
        st.markdown(
            f"<div style='text-align:center;margin-bottom:0.5rem;'>"
            f"<strong>Bienvenido/a</strong><br>"
            f"<span style='font-size:0.9rem;'>{usuario.get('nombre', usuario.get('username', ''))}</span><br>"
            f"<span style='font-size:0.78rem;opacity:0.8;'>Rol: {usuario.get('rol', '').capitalize()}</span>"
            f"</div>",
            unsafe_allow_html=True
        )
        
        st.markdown("---")
        
        # Menú de navegación
        opciones = {
            "📝 Nueva Entrevista":     "formulario",
            "🔑 Gestión de Casos CC":  "gestion_cc",
            "🔍 Revisión de Entrevistas": "revision",
            "📚 Recursos Normativos":  "recursos",
        }
        
        if usuario.get("rol") == "administrador":
            opciones["🗑️ Borrado Seguro"]      = "borrado"
            opciones["⚙️ Administración"]       = "administracion"
        
        for label, pagina in opciones.items():
            if st.button(label, use_container_width=True, key=f"nav_{pagina}"):
                st.session_state["pagina_actual"] = pagina
                st.rerun()
        
        st.markdown("---")
        
        # Recursos rápidos
        st.markdown(
            """
            <div style="text-align:center;">
                <a href="https://notebooklm.google.com/notebook/95aeeaf6-1f2a-4b4e-b7b1-b9d2ab3e10d9"
                   target="_blank"
                   style="color:#FAF3E0;font-size:0.85rem;text-decoration:none;">
                    📚 NotebookLM Normativo ↗
                </a>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        st.markdown("---")
        
        # Información del sistema
        df_cache = st.session_state.get("df_registro")
        n_regs   = len(df_cache) if df_cache is not None and not df_cache.empty else 0
        
        st.markdown(
            f"<div style='font-size:0.78rem;opacity:0.8;text-align:center;'>"
            f"📊 {n_regs} entrevista(s) registrada(s)<br>"
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')}"
            f"</div>",
            unsafe_allow_html=True
        )
        
        st.markdown("---")
        
        # Cerrar sesión
        if st.button("🚪 Cerrar Sesión", use_container_width=True, key="btn_logout"):
            cb_logout()
            st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# ENRUTADOR PRINCIPAL
# ─────────────────────────────────────────────────────────────────────────────

def main():
    """
    Función principal de la aplicación.
    Inicializa el estado, renderiza el sidebar y enruta a la página correcta.
    """
    # Inicializar base de datos y session_state
    inicializar_db()
    init_session_state()
    
    # Si no está autenticado, mostrar login
    if not st.session_state.get("autenticado"):
        pagina_login()
        return
    
    # Sidebar de navegación
    renderizar_sidebar()
    
    # Enrutamiento según página actual
    pagina = st.session_state.get("pagina_actual", "formulario")
    
    if pagina == "formulario":
        pagina_formulario()
    elif pagina == "gestion_cc":
        pagina_gestion_cc()
    elif pagina == "revision":
        pagina_revision()
    elif pagina == "borrado":
        pagina_borrado_seguro()
    elif pagina == "recursos":
        pagina_recursos()
    elif pagina == "administracion":
        pagina_administracion()
    else:
        pagina_formulario()


# ─────────────────────────────────────────────────────────────────────────────
# PUNTO DE ENTRADA
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    main()
